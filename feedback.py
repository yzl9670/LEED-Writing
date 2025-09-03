# feedback.py
<<<<<<< HEAD
# -*- coding: utf-8 -*-

"""
Overview
- No more chunking: index the student's whole narrative as a single document.
- Review order: PRIORITY items first (strict), then SUPPLEMENT items (feasibility/consistency).
- If neither `priority_items` nor `supplement_items` is provided, fall back to `leed_scores`
  (any key with a value > 0 is treated as an item to review).
- Returns: (feedback_text, scores_dict, err_msg). The scores_dict can be empty; it’s kept
  for front-end compatibility.

Environment
- Requires an OpenAI API key in `OPENAI_API_KEY`.
- Works with both the newer `OpenAI()` client and the legacy `openai.*.create` style SDK.
"""

from __future__ import annotations
import os
import io
import json
import uuid
import math
import logging
from typing import List, Dict, Any, Optional, Tuple

# ====== Logging ======
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

# ====== OpenAI client shim: support new and old SDKs ======
_OPENAI_MODE = None   # "new" | "old"
_client = None

def _init_openai():
    """Initialize OpenAI client for both new and legacy SDKs."""
    global _OPENAI_MODE, _client
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        logging.warning("OPENAI_API_KEY is not set. Please configure it in your environment.")
    try:
        # New SDK
        from openai import OpenAI  # type: ignore
        _client = OpenAI(api_key=api_key) if api_key else OpenAI()
        _OPENAI_MODE = "new"
        logging.info("Using NEW OpenAI SDK client.")
    except Exception:
        # Legacy SDK
        import openai  # type: ignore
        openai.api_key = api_key or None
        _client = openai
        _OPENAI_MODE = "old"
        logging.info("Using OLD OpenAI SDK (openai.*.create).")

_init_openai()


# ====== File reading helpers: .docx/.pdf/.txt ======
def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except Exception:
        logging.error("Missing dependency: python-docx")
        raise
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_pdf(path: str) -> str:
    try:
        import PyPDF2
    except Exception:
        logging.error("Missing dependency: PyPDF2")
        raise
    text = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                text.append(t)
    return "\n".join(text)


def read_text(user_input: Optional[str] = None, file_path: Optional[str] = None) -> str:
    """Return plain text from user input or an uploaded file."""
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return _read_docx(file_path)
        elif ext == ".pdf":
            return _read_pdf(file_path)
        else:
            with io.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    if user_input:
        return user_input.strip()
    return ""


# ====== OpenAI call wrappers ======
_EMBED_MODEL_CANDIDATES = ["text-embedding-3-small", "text-embedding-ada-002"]
_CHAT_MODEL_CANDIDATES = ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"]

def get_embedding(text: str) -> Optional[List[float]]:
    """Get an embedding vector for a string using the first working model."""
    text = (text or "").strip()
    if not text:
        return None

    for model in _EMBED_MODEL_CANDIDATES:
        try:
            if _OPENAI_MODE == "new":
                emb = _client.embeddings.create(model=model, input=text)  # type: ignore
                return emb.data[0].embedding  # type: ignore
            else:
                emb = _client.Embedding.create(model=model, input=text)  # type: ignore
                return emb["data"][0]["embedding"]
        except Exception as e:
            logging.warning(f"Embedding model '{model}' failed: {e}")
            continue
    logging.error("All embedding model candidates failed.")
    return None


def chat_once(system_prompt: str, user_prompt: str, temperature: float = 0.2, max_tokens: int = 300) -> str:
    """One-shot chat completion with fallback across model candidates."""
    last_err = None
    for model in _CHAT_MODEL_CANDIDATES:
        try:
            if _OPENAI_MODE == "new":
                resp = _client.chat.completions.create(  # type: ignore
                    model=model,
                    temperature=temperature,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=max_tokens
                )
                content = resp.choices[0].message.content or ""  # type: ignore
                return content.strip()
            else:
                resp = _client.ChatCompletion.create(  # type: ignore
                    model=model,
                    temperature=temperature,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=max_tokens
                )
                content = resp["choices"][0]["message"]["content"] or ""
                return content.strip()
        except Exception as e:
            last_err = e
            logging.warning(f"Chat model '{model}' failed: {e}")
            continue
    raise RuntimeError(f"All chat models failed. Last error: {last_err}")


# ====== Minimal local vector index (single full text) ======
class SimpleVectorIndex:
    """A tiny in-memory index for this request lifecycle."""
    def __init__(self):
        self._docs: List[str] = []
        self._embs: List[List[float]] = []
        self._metas: List[Dict[str, Any]] = []

    def add(self, document: str, embedding: List[float], metadata: Optional[Dict[str, Any]] = None):
        self._docs.append(document)
        self._embs.append(embedding)
        self._metas.append(metadata or {})

    @staticmethod
    def _cos_sim(a: List[float], b: List[float]) -> float:
        if not a or not b:
            return 0.0
        dot = 0.0
        na = 0.0
        nb = 0.0
        for x, y in zip(a, b):
            dot += x * y
            na += x * x
            nb += y * y
        if na <= 0 or nb <= 0:
            return 0.0
        return dot / (math.sqrt(na) * math.sqrt(nb))

    def query(self, query_embedding: List[float], n_results: int = 1) -> List[Dict[str, Any]]:
        if not self._docs:
            return []
        sims = [(i, self._cos_sim(query_embedding, emb)) for i, emb in enumerate(self._embs)]
        sims.sort(key=lambda x: x[1], reverse=True)
        out = []
        for i, score in sims[:max(1, n_results)]:
            out.append({
                "document": self._docs[i],
                "metadata": self._metas[i],
                "score": float(score)
            })
        return out


# Global index (per-process; repopulated per request)
_INDEX = SimpleVectorIndex()


def index_full_text(text: str) -> Optional[str]:
    """Index the entire student narrative as a single document."""
    try:
        emb = get_embedding(text)
        if not emb:
            logging.warning("Failed to embed the full text.")
            return None
        doc_id = str(uuid.uuid4())
        _INDEX.add(document=text, embedding=emb, metadata={"id": doc_id, "full_text": True})
        logging.debug("Full text added to index.")
        return doc_id
    except Exception as e:
        logging.exception(f"Error indexing full text: {e}")
        return None


# ====== Item-by-item review ======
def process_leed_item(item_name: str, importance: str = "priority") -> str:
    """
    Review a single LEED item.
    importance: 'priority' | 'supplement'
    """
    item_name = (item_name or "").strip()
    if not item_name:
        return "Empty item name."

    logging.debug(f"Processing item: {item_name} ({importance})")

    # We still run a trivial recall against the single full-text doc
    q = f"LEED item {importance}: {item_name}"
    q_emb = get_embedding(q)
    if not q_emb:
        return f"[{item_name}] Unable to create query embedding."

    results = _INDEX.query(query_embedding=q_emb, n_results=1)
    excerpt = results[0]["document"] if results else ""

    strict_note = (
        "This is a PRIORITY item. Be STRICT: prerequisites and critical thresholds must be explicitly satisfied. "
        "Identify any missing proofs, misinterpretations, or prerequisite gaps."
        if importance == "priority"
        else "This is a SUPPLEMENT item. Focus on plausibility, internal consistency, and whether the narrative proposes credible means of achievement."
    )

    prompt = f"""
You are an experienced LEED BD+C reviewer.
Focus on: "{item_name}".
{strict_note}

Student narrative (full-text excerpt; treat this as the student's entire narrative):
---
{excerpt}
---

Task:
1) State clearly whether the student's narrative sufficiently addresses "{item_name}".
2) If not fully sufficient, list the specific missing/unclear aspects.
3) Provide a concise, actionable recommendation (≤80 words).
Respond concisely, use bullet points if helpful, and avoid extra commentary.
"""

    try:
        resp = chat_once(
            system_prompt="You provide item-by-item LEED compliance feedback.",
            user_prompt=prompt,
            temperature=0.2,
            max_tokens=220
        )
        return resp.strip()
    except Exception as e:
        logging.exception(f"Chat error on item '{item_name}': {e}")
        return f"[{item_name}] Review failed: {e}"


# ====== Main entry point ======
def get_feedback(
    user_input: Optional[str] = None,
    file_path: Optional[str] = None,
    rubrics: Optional[Any] = None,  # kept for compatibility; not used
    leed_scores: Optional[Dict[str, Any]] = None,
    priority_items: Optional[List[Any]] = None,
    supplement_items: Optional[List[Any]] = None
) -> Tuple[str, Dict[str, Any], str]:
    """
    Read text → verify it's a LEED Narrative → index full text → review PRIORITY items first,
    then SUPPLEMENT items.

    - priority_items / supplement_items can be like:
        ["Optimize Energy Performance", ...] or [{"name": "Optimize Energy Performance"}, ...]
    - If both are empty, fall back to `leed_scores` (any key with value > 0 becomes an item).

    Returns: (final_feedback_text, scores_dict (pass-through), err_msg)
    """
    # 1) Read the narrative
    text = read_text(user_input=user_input, file_path=file_path)
    if not text or len(text.strip()) < 50:
        return "Your writing is too short to generate meaningful feedback.", {}, ""

    # 2) Quick classification: Is it a LEED Narrative?
    classify_prompt = f"""
Determine if the following text is specifically discussing a LEED Narrative for building certification.
If yes, respond with exactly "LEED Narrative".
If not, respond exactly "Not LEED Narrative".

Text:
{text[:6000]}  # truncated for safety
    """.strip()

    try:
        cls = chat_once(
            system_prompt="You classify whether the text is a LEED Narrative.",
            user_prompt=classify_prompt,
            temperature=0.0,
            max_tokens=5
        ).lower()
    except Exception as e:
        logging.exception(f"Classification error: {e}")
        return f"Error when checking narrative type: {e}", {}, ""

    if "leed narrative" not in cls:
        return "This passage is not a LEED Narrative.", {}, ""

    # 3) Index the full text
    index_full_text(text)

    # 4) Normalize priority/supplement arrays
    def _normalize_items(arr: Optional[List[Any]]) -> List[str]:
        if not arr:
            return []
        out = []
        for it in arr:
            if isinstance(it, dict) and "name" in it:
                out.append(str(it["name"]).strip())
            else:
                out.append(str(it).strip())
        return [x for x in out if x]

    priority_list = _normalize_items(priority_items)
    supplement_list = _normalize_items(supplement_items)

    # 5) Fallback: if both are empty, use leed_scores (keys with value > 0)
    fallback_list = []
    if not priority_list and not supplement_list and isinstance(leed_scores, dict):
        for k, v in leed_scores.items():
            if k == "total_score":
                continue
            try:
                val = float(v)
            except Exception:
                continue
            if val > 0:
                fallback_list.append(k)

    sections: List[str] = []

    # Priority items first (strict)
    if priority_list:
        pf = []
        for name in priority_list:
            pf.append(f"**{name}**\n{process_leed_item(name, importance='priority')}")
        sections.append("=== Priority Items Check ===\n" + "\n\n".join(pf))

    # Supplement items next (feasibility/consistency)
    if supplement_list:
        sf = []
        for name in supplement_list:
            sf.append(f"**{name}**\n{process_leed_item(name, importance='supplement')}")
        sections.append("=== Supplementary Items Check ===\n" + "\n\n".join(sf))

    # Fallback from numeric scores (treat as priority-level scrutiny)
    if not sections and fallback_list:
        items = []
        for name in fallback_list:
            items.append(f"**{name}**\n{process_leed_item(name, importance='priority')}")
        sections.append("=== Items From Scores (Fallback) ===\n" + "\n\n".join(items))

    final_feedback = "\n\n".join(sections) if sections else "No specific items were provided to review."
    # Currently we pass back `leed_scores` as-is for front-end compatibility.
    return final_feedback, (leed_scores or {}), ""


# ====== Local quick test ======
if __name__ == "__main__":
    demo_text = """
    Our project targets LEED v4 BD+C. We commit to meeting all prerequisites.
    We will optimize energy performance by improving the building envelope,
    adopting high-efficiency HVAC with VAV and heat recovery, and commissioning.
    Indoor water use reduction is addressed via WaterSense fixtures achieving >30% reduction baseline.
    We also consider enhanced refrigerant management by selecting low-GWP refrigerants.
    """
    fb, scores, err = get_feedback(
        user_input=demo_text,
        priority_items=["Optimize Energy Performance", "Indoor Water Use Reduction"],
        supplement_items=["Enhanced Refrigerant Management"]
    )
    print("=== FEEDBACK ===")
    print(fb)
    print("\n=== SCORES (passthrough) ===")
    print(json.dumps(scores, indent=2, ensure_ascii=False))
    if err:
        print("\nERR:", err)
=======

import os
import logging
import openai
import PyPDF2
from docx import Document
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
import chromadb
from chromadb.config import Settings

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("feedback.log"),
        logging.StreamHandler()
    ]
)


openai.api_key = os.getenv('OPENAI_API_KEY')

if not openai.api_key:
    logging.error("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
    raise EnvironmentError("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")

# Initialize ChromaDB, use PersistentClient and specify the persistence directory, and disable telemetry
persist_dir = "./temp_chroma"
if not os.path.exists(persist_dir):
    os.makedirs(persist_dir)
    logging.debug(f"Created persist_directory at {persist_dir}")

chroma_client = chromadb.PersistentClient(
    path=persist_dir,
    settings=Settings(anonymized_telemetry=False)  # Disable Telemetry
)
collection = chroma_client.get_or_create_collection(name="student_texts")

def chunk_text(full_text, chunk_size=400):
    """
    Chunking text for subsequent embedding generation and retrieval
    """
    logging.debug("Start chunking text.")
    if not full_text:
        logging.error("Text is empty, cannot chunk.")
        return []
    lines = []
    start = 0
    while start < len(full_text):
        end = start + chunk_size
        piece = full_text[start:end].strip()
        if piece:  # Make sure each piece is not empty
            lines.append(piece)
        start = end
    logging.debug(f"Total chunks created: {len(lines)}")
    return lines

def get_embeddings(texts):
    """Call OpenAI batch embedding generation interface"""
    try:
        resp = openai.Embedding.create(
            model="text-embedding-ada-002",
            input=texts
        )
        logging.debug(f"Generated embeddings for {len(texts)} texts.")
        return [item["embedding"] for item in resp["data"]]
    except openai.error.OpenAIError as e:
        logging.exception("Error generating embeddings:")
        return [None] * len(texts)

def get_embedding(text):
    """Calling OpenAI embedding generation interface"""
    try:
        resp = openai.Embedding.create(
            model="text-embedding-ada-002",
            input=text
        )
        logging.debug(f"Generated embedding for text: {text[:50]}...")
        return resp["data"][0]["embedding"]
    except openai.error.OpenAIError as e:
        logging.exception("Error generating embedding:")
        return None

def process_leed_item(item_name, collection):
    """
    Work on a single LEED project to generate feedback
    """
    logging.debug(f"Processing item: {item_name}")
    item_query = f"LEED item: {item_name}. Check compliance or missing info."
    try:
        item_emb = get_embedding(item_query)
        if not item_emb:
            return f"Failed to generate embedding for item: {item_name}"

        results = collection.query(
            query_embeddings=[item_emb],
            n_results=3,
        )
        relevant_docs = results.get("documents", [[]])[0]  # Check if the return value is empty

        if not relevant_docs:
            return f"No relevant documents found for item: {item_name}"

        context_text = "\n\n".join([f"Excerpt:\n{doc}" for doc in relevant_docs])
        prompt = f"""
        You are an experienced LEED BD+C reviewer.
        We focus on the item: {item_name}.

        Below are the most relevant excerpts from the student's text:
        {context_text}

        Task:
        1) Assess whether the student's text sufficiently addresses {item_name}.
        2) If the requirements are not fully met, list any missing or unclear aspects and provide a concise recommendation (under 100 words total).

        Respond strictly according to the instructions above, with no additional commentary or information.
        """
        gpt_resp = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You provide item-by-item LEED compliance feedback."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=150  # Adjust as needed
        )
        return gpt_resp["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logging.exception(f"Error processing item {item_name}:")
        return f"Error processing item {item_name}: {e}"

def process_leed_items(items, collection):
    """
    Work on multiple LEED projects, generating detailed feedback
    """
    logging.debug("Processing multiple LEED items.")
    detailed_feedback = []

    for item in items:
        item_name = item['name']
        feedback = process_leed_item(item_name, collection)
        if feedback != "This item is well addressed.":
            detailed_feedback.append(f"**{item_name}**:\n{feedback}")

    # Merge detailed feedback
    if detailed_feedback:
        detailed_feedback_str = "\n\n".join(detailed_feedback)
        final_feedback = f"=== Detailed LEED Item Feedback ===\n{detailed_feedback_str}"
    else:
        final_feedback = "All items are well addressed."

    logging.info("Feedback generation completed.")
    return final_feedback


def get_feedback(user_input=None, file_path=None, rubrics=None, leed_scores=None):
    """
    Core functionality: Get text from user input or file -> Classify -> Analyze LEED project -> Return feedback
    """

    # ========== 1. Reading input text ==========
    logging.debug("Start reading input text.")
    text = ""
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".docx":
                logging.debug(f"Reading DOCX file: {file_path}")
                doc = Document(file_path)
                text = "\n".join(para.text for para in doc.paragraphs)
            elif ext == ".pdf":
                logging.debug(f"Reading PDF file: {file_path}")
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            text += t
            else:
                logging.error(f"Unsupported file type: {ext}")
                return "Unsupported file type.", {}, ""
        except Exception as e:
            logging.exception("Error reading file:")
            return f"Error reading file: {e}", {}, ""
    elif user_input:
        logging.debug("Using user-provided input text.")
        text = user_input.strip()
    else:
        logging.error("No input provided.")
        return "No input provided.", {}, ""

    if len(text) < 50:
        logging.warning("Input text is too short.")
        return "Your writing is too short to get meaningful feedback.", {}, ""

    # ========== 2. Check if it is LEED Narrative ==========
    logging.debug("Classifying input as LEED Narrative or not.")
    classification_prompt = f"""
    Determine if the following text is specifically discussing a LEED Narrative for building certification. 
    If yes, respond with "LEED Narrative". 
    If not, respond "Not LEED Narrative".

    Text:
    {text}
    """
    try:
        cls_resp = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You classify if text is a LEED Narrative."},
                {"role": "user", "content": classification_prompt}
            ],
            temperature=0.0,
            timeout=10
        )
        cls_result = cls_resp["choices"][0]["message"]["content"].strip().lower()
        logging.debug(f"Classification result: {cls_result}")
    except openai.error.OpenAIError as e:
        logging.exception("Error calling OpenAI API:")
        return f"Error checking narrative type: {e}", {}, ""

    if "leed narrative" not in cls_result:
        logging.info("Input text is not related to LEED Narrative.")
        return "This passage is not related to LEED Narrative.", {}, ""

    logging.info("Input classified as LEED Narrative.")
    narrative_msg = "This passage is identified as part of a LEED Narrative."

    # ========== 3. Preparing for LEED Points ==========
    logging.debug("Preparing LEED scores.")
    item_dict = {}
    total_score = 0.0
    if leed_scores:
        try:
            total_score = float(leed_scores.get("total_score", 0))
            for k, v in leed_scores.items():
                if k == "total_score":
                    continue
                val = float(v)
                if val > 0:
                    item_dict[k] = val
        except (ValueError, TypeError) as e:
            logging.error(f"Invalid LEED scores provided: {e}")
            return f"Invalid LEED scores provided: {e}", {}, ""
    else:
        logging.info("No LEED scores provided.")
    
    logging.debug(f"LEED items to evaluate: {item_dict}")
    pass_msg = f"Total score is {total_score}. {'Pass (>=40)!' if total_score >= 40 else 'Less than 40, does not meet the LEED requirement.'}"

    # ========== 4. Block, embed and process LEED projects ==========
    logging.debug("Chunking text for analysis.")
    chunks = chunk_text(text, chunk_size=400)
    logging.debug(f"Total chunks created: {len(chunks)}")

    logging.debug("Embedding and storing text chunks in Chroma DB.")
    embeddings = get_embeddings(chunks)
    for i, (c, emb) in enumerate(zip(chunks, embeddings)):
        if not emb:
            logging.warning(f"Embedding failed for chunk {i}. Skipping...")
            continue
        doc_id = str(uuid.uuid4())
        try:
            collection.add(
                documents=[c],
                embeddings=[emb],
                ids=[doc_id],
                metadatas=[{"chunk_index": i}]
            )
            logging.debug(f"Successfully added chunk {i} to ChromaDB.")
        except Exception as e:
            logging.exception(f"Error embedding/storing chunk {i}: {c[:100]}")
            continue

    logging.debug("Finished embedding and storing all chunks in ChromaDB.")

    # Working on LEED projects
    logging.debug("Processing LEED items in parallel.")
    item_feedbacks = []
    if item_dict:
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {
                executor.submit(process_leed_item, item_name, collection): item_name
                for item_name in item_dict
            }

            for f in as_completed(futures):
                iname = futures[f]
                try:
                    result_text = f.result()
                    logging.debug(f"Feedback for {iname}: {result_text}")
                    if result_text != "This item is well addressed.":
                        item_feedbacks.append((iname, result_text))
                except Exception as e:
                    logging.exception(f"Error processing item {iname}:")
                    item_feedbacks.append((iname, f"Error processing item {iname}: {e}"))
    else:
        logging.info("No LEED items to process.")

    # Incorporate LEED Project Feedback
    if item_feedbacks:
        detailed_feedback_str = "\n\n".join([f"**{k}**:\n{v}" for k, v in item_feedbacks])
        final_feedback = f"=== Detailed LEED Item Feedback ===\n{detailed_feedback_str}"
    else:
        final_feedback = "All items are well addressed."

    logging.info("Feedback generation completed.")
    return final_feedback, leed_scores, ""

def read_pdf_in_chunks(file_path, chunk_size=1024):
    # 占位实现：你已经使用其它方法处理 PDF，此处返回空字符串
    return ""
>>>>>>> dac79ff818159051f5076784da316ba3e020095f
